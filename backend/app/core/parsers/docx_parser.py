"""
============================================================================
DOCX PARSER MODULE
============================================================================
This module handles extracting text from Microsoft Word (.docx) files.

KEY CONCEPTS TO UNDERSTAND:
1. DOCX Structure: DOCX files are actually ZIP archives containing XML
2. python-docx: Library that parses the XML structure
3. Paragraphs and Tables: Main content containers in Word documents
4. Formatting: Preserving bold, italic, bullets, etc.

INTERVIEW TALKING POINTS:
- "I use python-docx to parse Word documents, which handles the XML structure"
- "I extract both paragraphs and tables to capture all content"
- "The parser preserves formatting like bullets and numbering"

LIBRARY USED:
- python-docx: Official library for reading/writing Word documents
============================================================================
"""

from docx import Document
from pathlib import Path
from typing import Dict, Any, List
import logging
import re

# Set up logging
logger = logging.getLogger(__name__)


class DOCXParser:
    """
    Parse DOCX files and extract text content.
    
    HOW IT WORKS:
    1. Open DOCX file with python-docx
    2. Extract text from paragraphs
    3. Extract text from tables
    4. Preserve structure (bullets, numbering)
    5. Return organized content
    
    USAGE:
        parser = DOCXParser()
        result = parser.parse("resume.docx")
        print(result["text"])
    """
    
    def __init__(self):
        """Initialize the DOCX parser."""
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text.
        
        PARAMETERS:
            file_path: Path to the DOCX file
        
        RETURNS:
            Dictionary containing:
            - text: Extracted text content
            - paragraphs: Number of paragraphs
            - tables: Number of tables
            - metadata: Document properties
            - success: Whether extraction succeeded
        
        EXAMPLE:
            result = parser.parse("resume.docx")
            if result["success"]:
                print(result["text"])
        """
        file_path_obj = Path(file_path)
        
        # Validate file exists
        if not file_path_obj.exists():
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "text": "",
                "paragraphs": 0,
                "tables": 0,
                "metadata": {},
            }
        
        try:
            # Open the document
            doc = Document(file_path_obj)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract text from paragraphs
            paragraph_texts = self._extract_paragraphs(doc)
            
            # Extract text from tables
            table_texts = self._extract_tables(doc)
            
            # Combine all text
            all_text = []
            
            # Add paragraphs
            all_text.extend(paragraph_texts)
            
            # Add tables (if any)
            if table_texts:
                all_text.append("\n--- Tables ---\n")
                all_text.extend(table_texts)
            
            # Join and clean
            full_text = "\n".join(all_text)
            full_text = self._clean_text(full_text)
            
            self.logger.info(f"Successfully parsed {file_path}")
            
            return {
                "success": True,
                "text": full_text,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "metadata": metadata,
                "error": None
            }
        
        except Exception as e:
            self.logger.error(f"Failed to parse {file_path}: {e}")
            return {
                "success": False,
                "error": f"Failed to parse DOCX: {str(e)}",
                "text": "",
                "paragraphs": 0,
                "tables": 0,
                "metadata": {},
            }
    
    def _extract_metadata(self, doc: Document) -> Dict[str, str]:
        """
        Extract document metadata/properties.
        
        WHAT IT EXTRACTS:
        - Author
        - Title
        - Subject
        - Keywords
        - Last modified by
        - Created date
        
        WHY USEFUL:
        - Can help identify document owner
        - Useful for debugging
        - May contain additional context
        """
        metadata = {}
        
        try:
            core_props = doc.core_properties
            
            metadata = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "last_modified_by": core_props.last_modified_by or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            self.logger.warning(f"Failed to extract metadata: {e}")
        
        return metadata
    
    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """
        Extract text from all paragraphs.
        
        HOW IT WORKS:
        1. Iterate through all paragraphs
        2. Check for bullets/numbering
        3. Preserve structure with indentation
        4. Skip empty paragraphs
        
        FORMATTING PRESERVATION:
        - Bullets: Add "• " prefix
        - Numbering: Add number prefix
        - Indentation: Preserve with spaces
        """
        texts = []
        
        for para in doc.paragraphs:
            # Get paragraph text
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check if paragraph is a bullet point
            if para.style.name.startswith('List'):
                # Add bullet marker
                if 'Bullet' in para.style.name:
                    text = f"• {text}"
                # Numbered lists already have numbers in text
            
            texts.append(text)
        
        return texts
    
    def _extract_tables(self, doc: Document) -> List[str]:
        """
        Extract text from all tables.
        
        WHY IMPORTANT:
        - Resumes often use tables for layout
        - Skills, experience may be in tables
        - Need to preserve table structure
        
        HOW IT WORKS:
        1. Iterate through all tables
        2. Extract text from each cell
        3. Preserve row/column structure
        4. Format as readable text
        """
        table_texts = []
        
        for table_num, table in enumerate(doc.tables, 1):
            # Add table header
            table_texts.append(f"\n[Table {table_num}]")
            
            # Extract rows
            for row in table.rows:
                # Extract cells in this row
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                
                # Join cells with separator
                if cells:
                    row_text = " | ".join(cells)
                    table_texts.append(row_text)
        
        return table_texts
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        WHAT IT DOES:
        1. Remove excessive whitespace
        2. Normalize line breaks
        3. Remove special characters that cause issues
        4. Preserve structure
        """
        if not text:
            return ""
        
        # Normalize line breaks
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # Remove excessive blank lines (more than 2 consecutive)
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Remove excessive spaces (but keep single spaces)
        text = re.sub(r" {2,}", " ", text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        
        # Remove leading/trailing whitespace from entire text
        text = text.strip()
        
        return text
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Attempt to identify and extract resume sections.
        
        COMMON SECTIONS:
        - Contact Information
        - Summary/Objective
        - Work Experience
        - Education
        - Skills
        - Certifications
        
        NOTE: This is a basic implementation. The AI will do more
        sophisticated extraction.
        """
        sections = {}
        
        # Common section header patterns
        section_patterns = {
            "contact": r"(?i)(contact|personal information)",
            "summary": r"(?i)(summary|objective|profile|about)",
            "experience": r"(?i)(experience|employment|work history)",
            "education": r"(?i)(education|academic)",
            "skills": r"(?i)(skills|technical skills|competencies)",
            "certifications": r"(?i)(certifications?|licenses?)",
            "projects": r"(?i)(projects?)",
            "awards": r"(?i)(awards?|honors?|achievements?)",
        }
        
        # Split text into lines
        lines = text.split("\n")
        
        current_section = "header"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_header = False
            for section_name, pattern in section_patterns.items():
                if re.match(pattern, line):
                    # Save previous section
                    if current_content:
                        sections[current_section] = "\n".join(current_content)
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections


# ============================================================================
# CONVENIENCE FUNCTION
# ============================================================================

def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Convenience function to parse a DOCX file.
    
    USAGE:
        from app.core.parsers.docx_parser import parse_docx
        result = parse_docx("resume.docx")
        if result["success"]:
            print(result["text"])
    """
    parser = DOCXParser()
    return parser.parse(file_path)


# ============================================================================
# NOTES FOR INTERVIEWS
# ============================================================================
"""
KEY POINTS TO MENTION:

1. DOCX STRUCTURE:
   - DOCX is a ZIP file containing XML
   - python-docx parses the XML structure
   - More reliable than PDF parsing

2. CONTENT EXTRACTION:
   - Extract paragraphs (main text)
   - Extract tables (often used for layout)
   - Preserve formatting (bullets, numbering)

3. METADATA:
   - Extract document properties
   - Can identify author, creation date
   - Useful for validation

4. STRUCTURE PRESERVATION:
   - Maintain bullets and numbering
   - Preserve table structure
   - Keep section organization

5. ERROR HANDLING:
   - Handle corrupted files
   - Handle password-protected files
   - Graceful error messages

COMMON INTERVIEW QUESTIONS:

Q: "Why is DOCX easier to parse than PDF?"
A: "DOCX has a well-defined XML structure that's designed to be 
    machine-readable. PDFs are designed for display, not parsing, 
    so text extraction is more complex and error-prone."

Q: "How do you handle tables in resumes?"
A: "I extract tables separately and format them as readable text. 
    Many resumes use tables for layout, so this is crucial for 
    capturing all content. I preserve the row/column structure."

Q: "What about older .doc files?"
A: "This parser handles .docx (Office 2007+). For older .doc files, 
    I'd need to add support using libraries like antiword or convert 
    them to .docx first using tools like LibreOffice."

Q: "How would you improve this parser?"
A: "I could add:
    - Better formatting preservation (bold, italic, colors)
    - Image extraction (for logos, photos)
    - Hyperlink extraction
    - Style detection (fonts, sizes)
    - Support for older .doc format
    - Better table structure preservation"
"""

# Made with Bob
